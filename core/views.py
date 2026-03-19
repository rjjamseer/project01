import os
import tempfile
import uuid
import speech_recognition as sr
from django.shortcuts import render
from django.http import FileResponse
from docx import Document
import yt_dlp
from pydub import AudioSegment
import imageio_ffmpeg

# Explicitly set the FFmpeg path for pydub to avoid system PATH issues
AudioSegment.converter = imageio_ffmpeg.get_ffmpeg_exe()

def index(request):
    if request.method == 'POST':
        youtube_url = request.POST.get('url')
        if not youtube_url:
            return render(request, 'core/index.html', {'error': 'Please provide a valid YouTube URL.'})
        
        try:
            # Create a secure temp directory
            with tempfile.TemporaryDirectory() as temp_dir:
                # 1. Download Audio
                audio_path = os.path.join(temp_dir, 'audio_file')
                
                # yt-dlp config to get audio using the bundled ffmpeg binary
                ydl_opts = {
                    'format': 'worstaudio/worst',  # good enough for speech recognition, faster download
                    'outtmpl': audio_path + '.%(ext)s',
                    'ffmpeg_location': imageio_ffmpeg.get_ffmpeg_exe(),
                    'postprocessors': [{
                        'key': 'FFmpegExtractAudio',
                        'preferredcodec': 'wav',
                        'preferredquality': '128',
                    }],
                    'quiet': True,
                    'no_warnings': True,
                }
                
                with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                    info_dict = ydl.extract_info(youtube_url, download=True)
                    video_title = info_dict.get('title', 'Transcription')
                        
                wav_path = audio_path + '.wav'
                
                if not os.path.exists(wav_path):
                     return render(request, 'core/index.html', {'error': 'Failed to convert audio.'})
                
                # 2. Transcribe Audio
                recognizer = sr.Recognizer()
                
                # 1. Limit audio to the first 10 minutes for the demo
                audio = AudioSegment.from_wav(wav_path)
                max_duration_ms = 10 * 60 * 1000 # 10 mins
                demo_limit_reached = False
                if len(audio) > max_duration_ms:
                    audio = audio[:max_duration_ms]
                    demo_limit_reached = True
                
                # 2. Split into smaller 20-second chunks (Google Free API frequently times out on anything larger)
                chunk_length_ms = 20000 
                chunks = [audio[i:i+chunk_length_ms] for i in range(0, len(audio), chunk_length_ms)]
                
                full_text = []
                import time
                
                for i, chunk in enumerate(chunks):
                    chunk_path = os.path.join(temp_dir, f"chunk_{i}.wav")
                    chunk.export(chunk_path, format="wav")
                    
                    # Retry logic to prevent timeout failing the entire transcription
                    max_retries = 3
                    for attempt in range(max_retries):
                        with sr.AudioFile(chunk_path) as source:
                            audio_listened = recognizer.record(source)
                            try:
                                # Google Web Speech API (free)
                                text = recognizer.recognize_google(audio_listened)
                                full_text.append(text)
                                break # Success, break out of retry loop
                            except sr.UnknownValueError:
                                # Normal: Silence or unintelligible speech in this chunk
                                break
                            except sr.RequestError as e:
                                if attempt == max_retries - 1:
                                    full_text.append("[API Timeout for this segment]")
                                else:
                                    time.sleep(2) # Backoff before retrying
                
                if demo_limit_reached:
                    full_text.append("\n\n[Demo Limit Reached: Transcribed the first 10 minutes]")
                
                final_transcript = " ".join(full_text)
                if not final_transcript.strip():
                     final_transcript = "No perceptible speech found."
                
                # 3. Create Word Document
                doc = Document()
                doc.add_heading(f"{video_title}", 0)
                
                doc.add_heading('Transcript', level=1)
                doc.add_paragraph(final_transcript)
                
                doc.add_heading('My Notes', level=1)
                doc.add_paragraph()
                doc.add_paragraph("• ")
                doc.add_paragraph("• ")
                doc.add_paragraph("• ")
                
                # 4. Save and return file
                output_filename = f"{uuid.uuid4().hex[:8]}_notes.docx"
                output_filepath = os.path.join(tempfile.gettempdir(), output_filename)
                doc.save(output_filepath)
                
                return FileResponse(open(output_filepath, 'rb'), as_attachment=True, filename=f"Notes_Transcriber.docx")
                
        except Exception as e:
            return render(request, 'core/index.html', {'error': f"An error occurred: {str(e)}"})
            
    return render(request, 'core/index.html')
